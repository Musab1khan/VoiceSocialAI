import os
import logging
import magic
import cv2
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
import zipfile
import json
from typing import Dict, List, Optional, Tuple, Union
import base64
from io import BytesIO
from langdetect import detect

class FileProcessor:
    """
    Advanced file processor with OpenCV integration
    Handles PDF, Word, Excel, Images, ZIP files and camera captures
    """
    
    def __init__(self):
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'],
            'document': ['.pdf', '.doc', '.docx', '.txt'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'archive': ['.zip', '.rar', '.7z'],
            'video': ['.mp4', '.avi', '.mov', '.mkv'],
            'audio': ['.mp3', '.wav', '.ogg', '.flac']
        }
        
        # Create uploads directory
        self.upload_dir = "static/uploads"
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(f"{self.upload_dir}/processed", exist_ok=True)
        
    def process_file(self, file_data: bytes, filename: str) -> Dict:
        """
        Process uploaded file and extract information
        """
        try:
            # Save file temporarily
            file_path = os.path.join(self.upload_dir, filename)
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Detect file type
            file_type = self._detect_file_type(file_path)
            file_ext = os.path.splitext(filename)[1].lower()
            
            result = {
                'success': True,
                'filename': filename,
                'file_type': file_type,
                'file_size': len(file_data),
                'file_path': file_path,
                'processed_data': {}
            }
            
            # Process based on file type
            if file_type == 'image':
                result['processed_data'] = self._process_image(file_path)
            elif file_type == 'document':
                result['processed_data'] = self._process_document(file_path, file_ext)
            elif file_type == 'spreadsheet':
                result['processed_data'] = self._process_spreadsheet(file_path)
            elif file_type == 'archive':
                result['processed_data'] = self._process_archive(file_path)
            elif file_type == 'video':
                result['processed_data'] = self._process_video(file_path)
            else:
                result['processed_data'] = {'content': 'File type not supported for content extraction'}
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")
            return {
                'success': False,
                'error': str(e),
                'filename': filename
            }
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type using python-magic"""
        try:
            mime_type = magic.from_file(file_path, mime=True)
            
            if mime_type.startswith('image/'):
                return 'image'
            elif mime_type.startswith('video/'):
                return 'video'
            elif mime_type.startswith('audio/'):
                return 'audio'
            elif 'pdf' in mime_type:
                return 'document'
            elif 'word' in mime_type or 'document' in mime_type:
                return 'document'
            elif 'spreadsheet' in mime_type or 'excel' in mime_type:
                return 'spreadsheet'
            elif 'zip' in mime_type or 'archive' in mime_type:
                return 'archive'
            else:
                return 'unknown'
        except:
            # Fallback to file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            for file_type, extensions in self.supported_formats.items():
                if file_ext in extensions:
                    return file_type
            return 'unknown'
    
    def _process_image(self, file_path: str) -> Dict:
        """Process image with OpenCV analysis"""
        try:
            # Basic image info
            img = cv2.imread(file_path)
            height, width = img.shape[:2]
            
            # OpenCV analysis
            analysis = {
                'dimensions': f"{width}x{height}",
                'channels': img.shape[2] if len(img.shape) > 2 else 1,
                'size_mb': round(os.path.getsize(file_path) / (1024*1024), 2)
            }
            
            # Color analysis
            if len(img.shape) == 3:
                # Convert to different color spaces for analysis
                hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
                lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
                
                # Dominant colors
                analysis['dominant_colors'] = self._get_dominant_colors(img)
                
                # Brightness and contrast analysis
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                analysis['brightness'] = round(np.mean(gray), 2)
                analysis['contrast'] = round(np.std(gray), 2)
            
            # Face detection
            analysis['faces_detected'] = self._detect_faces(img)
            
            # Object detection (simplified)
            analysis['edges_detected'] = self._detect_edges(img)
            
            # Generate thumbnail
            thumbnail_path = self._generate_thumbnail(file_path)
            analysis['thumbnail'] = thumbnail_path
            
            return {
                'type': 'image_analysis',
                'analysis': analysis,
                'content': f"Image analysis complete: {analysis['dimensions']}, {analysis['faces_detected']} faces detected"
            }
            
        except Exception as e:
            logging.error(f"Error processing image: {e}")
            return {'content': f"Error analyzing image: {str(e)}"}
    
    def _process_document(self, file_path: str, file_ext: str) -> Dict:
        """Process document files (PDF, Word, etc.)"""
        try:
            content = ""
            metadata = {}
            
            if file_ext == '.pdf':
                reader = PdfReader(file_path)
                metadata['pages'] = len(reader.pages)
                
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                    
                if hasattr(reader, 'metadata') and reader.metadata:
                    metadata.update({
                        'title': reader.metadata.get('/Title', 'Unknown'),
                        'author': reader.metadata.get('/Author', 'Unknown'),
                        'subject': reader.metadata.get('/Subject', 'Unknown')
                    })
            
            elif file_ext in ['.docx', '.doc']:
                doc = Document(file_path)
                metadata['paragraphs'] = len(doc.paragraphs)
                
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
            
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            # Language detection
            if content.strip():
                try:
                    detected_language = detect(content[:1000])  # Use first 1000 chars
                    metadata['language'] = detected_language
                except:
                    metadata['language'] = 'unknown'
            
            # Content statistics
            words = len(content.split())
            chars = len(content)
            metadata.update({
                'word_count': words,
                'character_count': chars,
                'size_kb': round(os.path.getsize(file_path) / 1024, 2)
            })
            
            # Truncate content if too long
            preview = content[:2000] + "..." if len(content) > 2000 else content
            
            return {
                'type': 'document',
                'content': preview,
                'full_content': content,
                'metadata': metadata,
                'summary': f"Document with {words} words in {metadata.get('language', 'unknown')} language"
            }
            
        except Exception as e:
            logging.error(f"Error processing document: {e}")
            return {'content': f"Error reading document: {str(e)}"}
    
    def _process_spreadsheet(self, file_path: str) -> Dict:
        """Process spreadsheet files"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            
            metadata = {
                'sheets': workbook.sheetnames,
                'total_sheets': len(workbook.sheetnames)
            }
            
            # Process first sheet for preview
            first_sheet = workbook.active
            
            # Get data preview (first 10 rows, 5 columns)
            data_preview = []
            max_row = min(10, first_sheet.max_row)
            max_col = min(5, first_sheet.max_column)
            
            for row in range(1, max_row + 1):
                row_data = []
                for col in range(1, max_col + 1):
                    cell_value = first_sheet.cell(row=row, column=col).value
                    row_data.append(str(cell_value) if cell_value is not None else "")
                data_preview.append(row_data)
            
            metadata.update({
                'rows': first_sheet.max_row,
                'columns': first_sheet.max_column,
                'size_kb': round(os.path.getsize(file_path) / 1024, 2)
            })
            
            return {
                'type': 'spreadsheet',
                'content': f"Spreadsheet with {metadata['total_sheets']} sheets, {metadata['rows']} rows, {metadata['columns']} columns",
                'data_preview': data_preview,
                'metadata': metadata
            }
            
        except Exception as e:
            logging.error(f"Error processing spreadsheet: {e}")
            return {'content': f"Error reading spreadsheet: {str(e)}"}
    
    def _process_archive(self, file_path: str) -> Dict:
        """Process archive files (ZIP, etc.)"""
        try:
            if file_path.endswith('.zip'):
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    file_list = zip_file.namelist()
                    
                    metadata = {
                        'total_files': len(file_list),
                        'files': file_list[:20],  # Show first 20 files
                        'size_mb': round(os.path.getsize(file_path) / (1024*1024), 2)
                    }
                    
                    return {
                        'type': 'archive',
                        'content': f"ZIP archive containing {len(file_list)} files",
                        'metadata': metadata
                    }
            
            return {'content': 'Archive format not supported yet'}
            
        except Exception as e:
            logging.error(f"Error processing archive: {e}")
            return {'content': f"Error reading archive: {str(e)}"}
    
    def _process_video(self, file_path: str) -> Dict:
        """Process video files with OpenCV"""
        try:
            cap = cv2.VideoCapture(file_path)
            
            # Video metadata
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            metadata = {
                'duration_seconds': round(duration, 2),
                'fps': round(fps, 2),
                'resolution': f"{width}x{height}",
                'total_frames': frame_count,
                'size_mb': round(os.path.getsize(file_path) / (1024*1024), 2)
            }
            
            # Extract a few frames for preview
            frame_thumbnails = self._extract_video_frames(cap, 3)
            metadata['preview_frames'] = frame_thumbnails
            
            cap.release()
            
            return {
                'type': 'video',
                'content': f"Video: {metadata['resolution']}, {metadata['duration_seconds']}s, {metadata['fps']} FPS",
                'metadata': metadata
            }
            
        except Exception as e:
            logging.error(f"Error processing video: {e}")
            return {'content': f"Error analyzing video: {str(e)}"}
    
    def process_camera_capture(self, image_data: str) -> Dict:
        """Process camera captured image"""
        try:
            # Decode base64 image
            image_data = image_data.split(',')[1] if ',' in image_data else image_data
            image_bytes = base64.b64decode(image_data)
            
            # Save captured image
            import uuid
            filename = f"camera_capture_{uuid.uuid4().hex}.jpg"
            file_path = os.path.join(self.upload_dir, filename)
            
            with open(file_path, 'wb') as f:
                f.write(image_bytes)
            
            # Process the captured image
            result = self._process_image(file_path)
            result['source'] = 'camera_capture'
            result['filename'] = filename
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing camera capture: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _get_dominant_colors(self, img: np.ndarray, k: int = 5) -> List[List[int]]:
        """Extract dominant colors using K-means clustering"""
        try:
            # Reshape image to be a list of pixels
            data = img.reshape((-1, 3))
            data = np.float32(data)
            
            # Apply K-means
            criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 20, 1.0)
            _, labels, centers = cv2.kmeans(data, k, None, criteria, 10, cv2.KMEANS_RANDOM_CENTERS)
            
            # Convert to int and return
            centers = np.uint8(centers)
            return centers.tolist()
            
        except:
            return [[128, 128, 128]]  # Default gray if analysis fails
    
    def _detect_faces(self, img: np.ndarray) -> int:
        """Detect faces in image using OpenCV"""
        try:
            # Load face cascade classifier
            face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
            
            return len(faces)
            
        except:
            return 0
    
    def _detect_edges(self, img: np.ndarray) -> int:
        """Detect edges in image"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 100, 200)
            edge_count = np.count_nonzero(edges)
            
            return int(edge_count)
            
        except:
            return 0
    
    def _generate_thumbnail(self, file_path: str, size: Tuple[int, int] = (200, 200)) -> str:
        """Generate thumbnail for image"""
        try:
            img = Image.open(file_path)
            img.thumbnail(size, Image.Resampling.LANCZOS)
            
            # Save thumbnail
            filename = os.path.basename(file_path)
            name, ext = os.path.splitext(filename)
            thumbnail_filename = f"{name}_thumb{ext}"
            thumbnail_path = os.path.join(self.upload_dir, "processed", thumbnail_filename)
            
            img.save(thumbnail_path)
            return thumbnail_path
            
        except Exception as e:
            logging.error(f"Error generating thumbnail: {e}")
            return file_path
    
    def _extract_video_frames(self, cap: cv2.VideoCapture, num_frames: int = 3) -> List[str]:
        """Extract frames from video for preview"""
        try:
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            frame_paths = []
            
            for i in range(num_frames):
                frame_pos = int((i + 1) * frame_count / (num_frames + 1))
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_pos)
                
                ret, frame = cap.read()
                if ret:
                    # Save frame
                    import uuid
                    frame_filename = f"video_frame_{uuid.uuid4().hex}.jpg"
                    frame_path = os.path.join(self.upload_dir, "processed", frame_filename)
                    cv2.imwrite(frame_path, frame)
                    frame_paths.append(frame_path)
            
            return frame_paths
            
        except Exception as e:
            logging.error(f"Error extracting video frames: {e}")
            return []

# Create global instance
file_processor = FileProcessor()